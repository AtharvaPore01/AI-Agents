from uagents import Agent, Context, Model
import os
import docx
import PyPDF2
import fpdf
from reportlab.pdfgen import canvas
import pdfplumber
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Define request and response models
class ResumeFile(Model):
    file_path: str

class ProcessedResume(Model):
    text: str
    output_file_path: str

# Define the Preprocessing Agent
preprocessing_agent = Agent(
    name="Preprocessing Agent",
    port=4041,
    endpoint=["http://localhost:4041/submit"],
    seed="preprocessing_server_seed"
)

# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text.strip()

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

# Function to preprocess text (cleaning extra spaces)
def preprocess_text(text):
    return " ".join(text.split())

def save_text_with_reportlab(text, output_file_path):
    c = canvas.Canvas(output_file_path, pagesize=letter)
    width, height = letter
    
    y_position = height - 40  # Start position
    text_object = c.beginText(40, y_position)
    text_object.setFont("Helvetica", 12)
    
    # Add text line by line
    for line in text.split('\n'):
        text_object.textLine(line)
    
    c.drawText(text_object)
    c.save()


# Function to save processed text as DOCX
def save_text_as_docx(text, output_path):
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(output_path)

@preprocessing_agent.on_event("startup")
async def startup_handler(ctx: Context):
    ctx.logger.info("Preprocessing Agent is ready.")
    ctx.logger.info("Waiting for a file path from the client...")

@preprocessing_agent.on_message(model=ResumeFile)
async def handle_resume(ctx: Context, sender: str, msg: ResumeFile):  # <-- Added sender
    ctx.logger.info(f"Received resume file path from {sender}: {msg.file_path}")

    if not os.path.exists(msg.file_path):
        ctx.logger.error("File not found. Please check the path.")
        return
    
    file_extension = os.path.splitext(msg.file_path)[-1].lower()
    extracted_text = ""

    if file_extension == ".pdf":
        extracted_text = extract_text_from_pdf(msg.file_path)
    elif file_extension == ".docx":
        extracted_text = extract_text_from_docx(msg.file_path)
    else:
        ctx.logger.error("Unsupported file format.")
        return

    print("------------extracted text-----------",extracted_text)

    # Determine output file format
    output_file_path = msg.file_path.replace(file_extension, "_processed" + file_extension)

    if file_extension == ".pdf":
        save_text_with_reportlab(extracted_text, output_file_path)
    elif file_extension == ".docx":
        save_text_with_reportlab(extracted_text, output_file_path)

    ctx.logger.info(f"Resume preprocessing complete. Saved at: {output_file_path}")

    # Send processed resume back to the sender
    await ctx.send(sender, ProcessedResume(text=extracted_text, output_file_path=output_file_path))

if __name__ == "__main__":
    preprocessing_agent.run()
